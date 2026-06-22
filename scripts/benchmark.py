"""Comprehensive NavyFox vs python-docx performance comparison.

Benchmark categories:
  lifecycle     — document create and open/close overhead
  write_paras   — build N plain paragraphs + save
  write_paras_fmt — N paragraphs with style/alignment/spacing + save
  write_runs    — single paragraph with N plain runs + save
  write_runs_fmt  — N runs with bold/italic/color/font_size + save
  write_table   — N-row x 3-col table + save
  write_mixed   — N paragraphs + N/10 table rows + save  (original benchmark)
  read_paras    — open doc, read all para.text
  read_para_props — open doc, read para style/alignment/spacing
  read_runs     — open doc, iterate runs, read run.text
  read_run_props  — open doc, read run bold/italic/color/font_size
  read_table    — open doc, read all table cell.text
  read_mixed    — open doc, read all paragraphs + all table cells (original benchmark)

Results are written to scripts/benchmark_results.csv.

Usage:
    python scripts/benchmark.py           # full run (n up to 10 000)
    python scripts/benchmark.py --quick   # fast sanity check (n up to 1 000)

Note: table benchmarks (write_table, read_table) are capped at TABLE_SIZES because
python-docx table creation is O(n²) and takes ~75 s at n=1 000, making larger sizes
impractical.  Mixed-document benchmarks use n/10 rows and scale normally.
"""

from __future__ import annotations

import argparse
import contextlib
import csv
import os
import statistics
import tempfile
import time
from collections.abc import Callable

SIZES = [100, 500, 1_000, 5_000, 10_000]
SIZES_QUICK = [100, 500, 1_000]

# python-docx table write is O(n²) via XML; cap at 1 000 rows to avoid hour-long runs.
TABLE_SIZES = [100, 500, 1_000]
TABLE_SIZES_QUICK = [50, 100]

# lifecycle benchmarks have fixed N=1
LIFECYCLE_REPEATS = 20


def _repeats(n: int) -> int:
    if n >= 5_000:
        return 2
    if n >= 1_000:
        return 3
    return 5


def _temp_path(suffix: str = ".docx") -> str:
    fd, path = tempfile.mkstemp(suffix=suffix)
    os.close(fd)
    return path


def _measure(fn: Callable[[], None], repeats: int) -> list[float]:
    """Return raw elapsed seconds for each of *repeats* calls of fn()."""
    times: list[float] = []
    for _ in range(repeats):
        t0 = time.perf_counter()
        fn()
        times.append(time.perf_counter() - t0)
    return times


# ---------------------------------------------------------------------------
# Reference-file factory (shared between read benchmarks)
# ---------------------------------------------------------------------------


def _build_ref_plain_paras(n: int) -> str:
    """Create a .docx with N plain paragraphs using python-docx."""
    import docx

    path = _temp_path()
    doc = docx.Document()
    for i in range(n):
        doc.add_paragraph(f"Paragraph {i}: The quick brown fox jumps over the lazy dog.")
    doc.save(path)
    return path


def _build_ref_formatted_paras(n: int) -> str:
    """Create a .docx with N paragraphs with inline formatting using python-docx."""
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    path = _temp_path()
    doc = docx.Document()
    for i in range(n):
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER if i % 2 == 0 else WD_ALIGN_PARAGRAPH.LEFT
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        run = para.add_run(f"Para {i}: formatted text here.")
        run.bold = i % 3 == 0
        run.italic = i % 4 == 0
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 128)
    doc.save(path)
    return path


def _build_ref_runs(n: int) -> str:
    """Create a .docx with one paragraph containing N runs using python-docx."""
    import docx

    path = _temp_path()
    doc = docx.Document()
    para = doc.add_paragraph()
    for i in range(n):
        para.add_run(f"Run {i}. ")
    doc.save(path)
    return path


def _build_ref_runs_formatted(n: int) -> str:
    """Create a .docx with one paragraph containing N formatted runs."""
    import docx
    from docx.shared import Pt, RGBColor

    path = _temp_path()
    doc = docx.Document()
    para = doc.add_paragraph()
    for i in range(n):
        run = para.add_run(f"Run {i}. ")
        run.bold = i % 2 == 0
        run.italic = i % 3 == 0
        run.font.size = Pt(11 + i % 4)
        run.font.color.rgb = RGBColor(0, 0, 128)
    doc.save(path)
    return path


def _build_ref_table(n: int) -> str:
    """Create a .docx with a single N-row × 3-col table using python-docx."""
    import docx

    path = _temp_path()
    doc = docx.Document()
    table = doc.add_table(rows=n, cols=3)
    for r in range(n):
        table.cell(r, 0).text = f"Row {r}"
        table.cell(r, 1).text = "Value A"
        table.cell(r, 2).text = "Value B"
    doc.save(path)
    return path


def _build_ref_mixed(n: int) -> str:
    """Create a .docx with N paragraphs + N/10 table rows using python-docx."""
    import docx

    path = _temp_path()
    doc = docx.Document()
    for i in range(n):
        doc.add_paragraph(f"Paragraph {i}: The quick brown fox jumps over the lazy dog.")
    n_rows = max(1, n // 10)
    table = doc.add_table(rows=n_rows, cols=3)
    for r in range(n_rows):
        table.cell(r, 0).text = f"Row {r}"
        table.cell(r, 1).text = "Value A"
        table.cell(r, 2).text = "Value B"
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# Lifecycle benchmarks
# ---------------------------------------------------------------------------


def bench_lifecycle_create_docx() -> list[float]:
    def fn():
        import docx

        doc = docx.Document()
        path = _temp_path()
        try:
            doc.save(path)
        finally:
            os.unlink(path)

    return _measure(fn, LIFECYCLE_REPEATS)


def bench_lifecycle_create_navyfox() -> list[float]:
    def fn():
        from navyfox import Document

        path = _temp_path()
        try:
            with Document() as doc:
                doc.save(path)
        finally:
            os.unlink(path)

    return _measure(fn, LIFECYCLE_REPEATS)


def bench_lifecycle_open_close_docx(ref_path: str) -> list[float]:
    def fn():
        import docx

        docx.Document(ref_path)  # python-docx has no explicit close

    return _measure(fn, LIFECYCLE_REPEATS)


def bench_lifecycle_open_close_navyfox(ref_path: str) -> list[float]:
    def fn():
        from navyfox import Document

        doc = Document.open(ref_path)
        doc.close()

    return _measure(fn, LIFECYCLE_REPEATS)


# ---------------------------------------------------------------------------
# Write: plain paragraphs
# ---------------------------------------------------------------------------


def bench_write_paras_docx(n: int) -> list[float]:
    import docx

    def fn():
        path = _temp_path()
        try:
            doc = docx.Document()
            for i in range(n):
                doc.add_paragraph(f"Paragraph {i}: The quick brown fox jumps over the lazy dog.")
            doc.save(path)
        finally:
            os.unlink(path)

    return _measure(fn, _repeats(n))


def bench_write_paras_navyfox(n: int) -> list[float]:
    from navyfox import Document

    def fn():
        path = _temp_path()
        try:
            with Document() as doc:
                for i in range(n):
                    doc.add_paragraph(
                        f"Paragraph {i}: The quick brown fox jumps over the lazy dog."
                    )
                doc.save(path)
        finally:
            os.unlink(path)

    return _measure(fn, _repeats(n))


# ---------------------------------------------------------------------------
# Write: formatted paragraphs (style, alignment, spacing)
# ---------------------------------------------------------------------------


def bench_write_paras_fmt_docx(n: int) -> list[float]:
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    def fn():
        path = _temp_path()
        try:
            doc = docx.Document()
            for i in range(n):
                para = doc.add_paragraph()
                pf = para.paragraph_format
                pf.alignment = WD_ALIGN_PARAGRAPH.CENTER if i % 2 == 0 else WD_ALIGN_PARAGRAPH.LEFT
                pf.space_before = Pt(6)
                pf.space_after = Pt(6)
                para.add_run(f"Para {i}: formatted text.")
            doc.save(path)
        finally:
            os.unlink(path)

    return _measure(fn, _repeats(n))


def bench_write_paras_fmt_navyfox(n: int) -> list[float]:
    from navyfox import Document, Paragraph

    def fn():
        path = _temp_path()
        try:
            with Document() as doc:
                for i in range(n):
                    doc.paragraphs.append(
                        Paragraph(
                            f"Para {i}: formatted text.",
                            alignment="center" if i % 2 == 0 else "left",
                            space_before=6.0,
                            space_after=6.0,
                        )
                    )
                doc.save(path)
        finally:
            os.unlink(path)

    return _measure(fn, _repeats(n))


# ---------------------------------------------------------------------------
# Write: runs (plain)
# ---------------------------------------------------------------------------


def bench_write_runs_docx(n: int) -> list[float]:
    import docx

    def fn():
        path = _temp_path()
        try:
            doc = docx.Document()
            para = doc.add_paragraph()
            for i in range(n):
                para.add_run(f"Run {i}. ")
            doc.save(path)
        finally:
            os.unlink(path)

    return _measure(fn, _repeats(n))


def bench_write_runs_navyfox(n: int) -> list[float]:
    from navyfox import Document, Run

    def fn():
        path = _temp_path()
        try:
            with Document() as doc:
                para = doc.add_paragraph()
                for i in range(n):
                    para.runs.append(Run(f"Run {i}. "))
                doc.save(path)
        finally:
            os.unlink(path)

    return _measure(fn, _repeats(n))


# ---------------------------------------------------------------------------
# Write: formatted runs (bold, italic, color, font_size)
# ---------------------------------------------------------------------------


def bench_write_runs_fmt_docx(n: int) -> list[float]:
    import docx
    from docx.shared import Pt, RGBColor

    def fn():
        path = _temp_path()
        try:
            doc = docx.Document()
            para = doc.add_paragraph()
            for i in range(n):
                run = para.add_run(f"Run {i}. ")
                run.bold = i % 2 == 0
                run.italic = i % 3 == 0
                run.font.size = Pt(11 + i % 4)
                run.font.color.rgb = RGBColor(0, 0, 128)
            doc.save(path)
        finally:
            os.unlink(path)

    return _measure(fn, _repeats(n))


def bench_write_runs_fmt_navyfox(n: int) -> list[float]:
    from navyfox import Document, Run

    def fn():
        path = _temp_path()
        try:
            with Document() as doc:
                para = doc.add_paragraph()
                for i in range(n):
                    para.runs.append(
                        Run(
                            f"Run {i}. ",
                            bold=i % 2 == 0,
                            italic=i % 3 == 0,
                            font_size=float(11 + i % 4),
                            color="#00008B",
                        )
                    )
                doc.save(path)
        finally:
            os.unlink(path)

    return _measure(fn, _repeats(n))


# ---------------------------------------------------------------------------
# Write: table
# ---------------------------------------------------------------------------


def bench_write_table_docx(n: int) -> list[float]:
    import docx

    def fn():
        path = _temp_path()
        try:
            doc = docx.Document()
            table = doc.add_table(rows=n, cols=3)
            for r in range(n):
                table.cell(r, 0).text = f"Row {r}"
                table.cell(r, 1).text = "Value A"
                table.cell(r, 2).text = "Value B"
            doc.save(path)
        finally:
            os.unlink(path)

    return _measure(fn, _repeats(n))


def bench_write_table_navyfox(n: int) -> list[float]:
    from navyfox import Document

    def fn():
        path = _temp_path()
        try:
            with Document() as doc:
                table = doc.add_table(rows=n, cols=3)
                for r in range(n):
                    table[r, 0].text = f"Row {r}"
                    table[r, 1].text = "Value A"
                    table[r, 2].text = "Value B"
                doc.save(path)
        finally:
            os.unlink(path)

    return _measure(fn, _repeats(n))


# ---------------------------------------------------------------------------
# Write: mixed (original benchmark)
# ---------------------------------------------------------------------------


def bench_write_mixed_docx(n: int) -> list[float]:
    import docx

    def fn():
        path = _temp_path()
        try:
            doc = docx.Document()
            for i in range(n):
                doc.add_paragraph(f"Paragraph {i}: The quick brown fox jumps over the lazy dog.")
            n_rows = max(1, n // 10)
            table = doc.add_table(rows=n_rows, cols=3)
            for r in range(n_rows):
                table.cell(r, 0).text = f"Row {r}"
                table.cell(r, 1).text = "Value A"
                table.cell(r, 2).text = "Value B"
            doc.save(path)
        finally:
            os.unlink(path)

    return _measure(fn, _repeats(n))


def bench_write_mixed_navyfox(n: int) -> list[float]:
    from navyfox import Document

    def fn():
        path = _temp_path()
        try:
            with Document() as doc:
                for i in range(n):
                    doc.add_paragraph(
                        f"Paragraph {i}: The quick brown fox jumps over the lazy dog."
                    )
                n_rows = max(1, n // 10)
                table = doc.add_table(rows=n_rows, cols=3)
                for r in range(n_rows):
                    table[r, 0].text = f"Row {r}"
                    table[r, 1].text = "Value A"
                    table[r, 2].text = "Value B"
                doc.save(path)
        finally:
            os.unlink(path)

    return _measure(fn, _repeats(n))


# ---------------------------------------------------------------------------
# Read: plain paragraph text
# ---------------------------------------------------------------------------


def bench_read_paras_docx(path: str, n: int) -> list[float]:
    import docx

    def fn():
        doc = docx.Document(path)
        for p in doc.paragraphs:
            _ = p.text

    return _measure(fn, _repeats(n))


def bench_read_paras_navyfox(path: str, n: int) -> list[float]:
    from navyfox import Document

    def fn():
        with Document.open(path) as doc:
            for p in doc.paragraphs:
                _ = p.text

    return _measure(fn, _repeats(n))


# ---------------------------------------------------------------------------
# Read: paragraph properties (style, alignment, spacing)
# ---------------------------------------------------------------------------


def bench_read_para_props_docx(path: str, n: int) -> list[float]:
    import docx

    def fn():
        doc = docx.Document(path)
        for p in doc.paragraphs:
            _ = p.text
            _ = p.style.name if p.style else None
            _ = p.alignment
            _ = p.paragraph_format.space_before
            _ = p.paragraph_format.space_after

    return _measure(fn, _repeats(n))


def bench_read_para_props_navyfox(path: str, n: int) -> list[float]:
    from navyfox import Document

    def fn():
        with Document.open(path) as doc:
            for p in doc.paragraphs:
                _ = p.text
                _ = p.style
                _ = p.alignment
                _ = p.space_before
                _ = p.space_after

    return _measure(fn, _repeats(n))


# ---------------------------------------------------------------------------
# Read: run text
# ---------------------------------------------------------------------------


def bench_read_runs_docx(path: str, n: int) -> list[float]:
    import docx

    def fn():
        doc = docx.Document(path)
        for p in doc.paragraphs:
            for run in p.runs:
                _ = run.text

    return _measure(fn, _repeats(n))


def bench_read_runs_navyfox(path: str, n: int) -> list[float]:
    from navyfox import Document

    def fn():
        with Document.open(path) as doc:
            for p in doc.paragraphs:
                for run in p.runs:
                    _ = run.text

    return _measure(fn, _repeats(n))


# ---------------------------------------------------------------------------
# Read: run formatting properties
# ---------------------------------------------------------------------------


def bench_read_run_props_docx(path: str, n: int) -> list[float]:
    import docx

    def fn():
        doc = docx.Document(path)
        for p in doc.paragraphs:
            for run in p.runs:
                _ = run.text
                _ = run.bold
                _ = run.italic
                _ = run.font.size
                with contextlib.suppress(Exception):
                    _ = run.font.color.rgb

    return _measure(fn, _repeats(n))


def bench_read_run_props_navyfox(path: str, n: int) -> list[float]:
    from navyfox import Document

    def fn():
        with Document.open(path) as doc:
            for p in doc.paragraphs:
                for run in p.runs:
                    _ = run.text
                    _ = run.bold
                    _ = run.italic
                    _ = run.font_size
                    _ = run.color

    return _measure(fn, _repeats(n))


# ---------------------------------------------------------------------------
# Read: table cells
# ---------------------------------------------------------------------------


def bench_read_table_docx(path: str, n: int) -> list[float]:
    import docx

    def fn():
        doc = docx.Document(path)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    _ = cell.text

    return _measure(fn, _repeats(n))


def bench_read_table_navyfox(path: str, n: int) -> list[float]:
    from navyfox import Document

    def fn():
        with Document.open(path) as doc:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        _ = cell.text

    return _measure(fn, _repeats(n))


# ---------------------------------------------------------------------------
# Read: mixed (original benchmark)
# ---------------------------------------------------------------------------


def bench_read_mixed_docx(path: str, n: int) -> list[float]:
    import docx

    def fn():
        doc = docx.Document(path)
        for p in doc.paragraphs:
            _ = p.text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    _ = cell.text

    return _measure(fn, _repeats(n))


def bench_read_mixed_navyfox(path: str, n: int) -> list[float]:

    from navyfox import Document

    def fn():
        with Document.open(path) as doc:
            for p in doc.paragraphs:
                _ = p.text
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        _ = cell.text

    return _measure(fn, _repeats(n))


# ---------------------------------------------------------------------------
# Result accumulation
# ---------------------------------------------------------------------------

_rows: list[dict[str, object]] = []


def _record(
    category: str,
    operation: str,
    library: str,
    n: int,
    times: list[float],
) -> None:
    for i, t in enumerate(times):
        _rows.append(
            {
                "category": category,
                "operation": operation,
                "library": library,
                "n": n,
                "repeat": i,
                "elapsed_s": round(t, 9),
                "elapsed_ms": round(t * 1000, 6),
            }
        )


def _run(
    label: str,
    category: str,
    operation: str,
    library: str,
    n: int,
    fn: Callable[..., list[float]],
    *args: object,
) -> None:
    print(f"  {label:<50}", end=" ", flush=True)
    times = fn(*args)
    _record(category, operation, library, n, times)
    mean = statistics.mean(times)
    stdev = statistics.stdev(times) if len(times) > 1 else 0.0
    print(f"{mean * 1000:8.2f} ms  ±{stdev * 1000:.2f}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def main(sizes: list[int], table_sizes: list[int]) -> None:
    # ---- lifecycle --------------------------------------------------------
    print("\n=== Lifecycle (empty document create + save) ===")
    print(f"  {'python-docx':<50}", end=" ", flush=True)
    times = bench_lifecycle_create_docx()
    _record("lifecycle", "create_empty", "python_docx", 1, times)
    mean = statistics.mean(times)
    print(f"{mean * 1000:8.2f} ms  ±{statistics.stdev(times) * 1000:.2f}")

    print(f"  {'navyfox':<50}", end=" ", flush=True)
    times = bench_lifecycle_create_navyfox()
    _record("lifecycle", "create_empty", "navyfox", 1, times)
    mean = statistics.mean(times)
    print(f"{mean * 1000:8.2f} ms  ±{statistics.stdev(times) * 1000:.2f}")

    print("\n=== Lifecycle (open + close existing doc) ===")
    _ref_small = _build_ref_plain_paras(10)
    try:
        print(f"  {'python-docx':<50}", end=" ", flush=True)
        times = bench_lifecycle_open_close_docx(_ref_small)
        _record("lifecycle", "open_close", "python_docx", 1, times)
        mean = statistics.mean(times)
        print(f"{mean * 1000:8.2f} ms  ±{statistics.stdev(times) * 1000:.2f}")

        print(f"  {'navyfox':<50}", end=" ", flush=True)
        times = bench_lifecycle_open_close_navyfox(_ref_small)
        _record("lifecycle", "open_close", "navyfox", 1, times)
        mean = statistics.mean(times)
        print(f"{mean * 1000:8.2f} ms  ±{statistics.stdev(times) * 1000:.2f}")
    finally:
        os.unlink(_ref_small)

    # ---- write: plain paragraphs ------------------------------------------
    print("\n=== Write: plain paragraphs ===")
    for n in sizes:
        _run(
            f"python-docx  n={n:>6}",
            "write",
            "write_paras",
            "python_docx",
            n,
            bench_write_paras_docx,
            n,
        )
        _run(
            f"navyfox      n={n:>6}",
            "write",
            "write_paras",
            "navyfox",
            n,
            bench_write_paras_navyfox,
            n,
        )

    # ---- write: formatted paragraphs --------------------------------------
    print("\n=== Write: formatted paragraphs (style/alignment/spacing) ===")
    for n in sizes:
        _run(
            f"python-docx  n={n:>6}",
            "write",
            "write_paras_fmt",
            "python_docx",
            n,
            bench_write_paras_fmt_docx,
            n,
        )
        _run(
            f"navyfox      n={n:>6}",
            "write",
            "write_paras_fmt",
            "navyfox",
            n,
            bench_write_paras_fmt_navyfox,
            n,
        )

    # ---- write: plain runs ------------------------------------------------
    print("\n=== Write: plain runs (N runs into one paragraph) ===")
    for n in sizes:
        _run(
            f"python-docx  n={n:>6}",
            "write",
            "write_runs",
            "python_docx",
            n,
            bench_write_runs_docx,
            n,
        )
        _run(
            f"navyfox      n={n:>6}",
            "write",
            "write_runs",
            "navyfox",
            n,
            bench_write_runs_navyfox,
            n,
        )

    # ---- write: formatted runs --------------------------------------------
    print("\n=== Write: formatted runs (bold/italic/color/font_size) ===")
    for n in sizes:
        _run(
            f"python-docx  n={n:>6}",
            "write",
            "write_runs_fmt",
            "python_docx",
            n,
            bench_write_runs_fmt_docx,
            n,
        )
        _run(
            f"navyfox      n={n:>6}",
            "write",
            "write_runs_fmt",
            "navyfox",
            n,
            bench_write_runs_fmt_navyfox,
            n,
        )

    # ---- write: table -----------------------------------------------------
    print("\n=== Write: table (N rows × 3 cols) ===")
    print(f"  [capped at max n={table_sizes[-1]}; python-docx table creation is O(n²)]")
    for n in table_sizes:
        _run(
            f"python-docx  n={n:>6}",
            "write",
            "write_table",
            "python_docx",
            n,
            bench_write_table_docx,
            n,
        )
        _run(
            f"navyfox      n={n:>6}",
            "write",
            "write_table",
            "navyfox",
            n,
            bench_write_table_navyfox,
            n,
        )

    # ---- write: mixed -----------------------------------------------------
    print("\n=== Write: mixed (N paragraphs + N/10 table rows) ===")
    for n in sizes:
        _run(
            f"python-docx  n={n:>6}",
            "write",
            "write_mixed",
            "python_docx",
            n,
            bench_write_mixed_docx,
            n,
        )
        _run(
            f"navyfox      n={n:>6}",
            "write",
            "write_mixed",
            "navyfox",
            n,
            bench_write_mixed_navyfox,
            n,
        )

    # ---- read: plain paragraph text ---------------------------------------
    print("\n=== Read: plain paragraph text ===")
    for n in sizes:
        ref = _build_ref_plain_paras(n)
        print(f"  [ref n={n}]")
        try:
            _run(
                f"python-docx  n={n:>6}",
                "read",
                "read_paras",
                "python_docx",
                n,
                bench_read_paras_docx,
                ref,
                n,
            )
            _run(
                f"navyfox      n={n:>6}",
                "read",
                "read_paras",
                "navyfox",
                n,
                bench_read_paras_navyfox,
                ref,
                n,
            )
        finally:
            os.unlink(ref)

    # ---- read: paragraph properties ---------------------------------------
    print("\n=== Read: paragraph properties (text + style/alignment/spacing) ===")
    for n in sizes:
        ref = _build_ref_formatted_paras(n)
        print(f"  [ref n={n}]")
        try:
            _run(
                f"python-docx  n={n:>6}",
                "read",
                "read_para_props",
                "python_docx",
                n,
                bench_read_para_props_docx,
                ref,
                n,
            )
            _run(
                f"navyfox      n={n:>6}",
                "read",
                "read_para_props",
                "navyfox",
                n,
                bench_read_para_props_navyfox,
                ref,
                n,
            )
        finally:
            os.unlink(ref)

    # ---- read: run text ---------------------------------------------------
    print("\n=== Read: run text ===")
    for n in sizes:
        ref = _build_ref_runs(n)
        print(f"  [ref n={n} runs in 1 para]")
        try:
            _run(
                f"python-docx  n={n:>6}",
                "read",
                "read_runs",
                "python_docx",
                n,
                bench_read_runs_docx,
                ref,
                n,
            )
            _run(
                f"navyfox      n={n:>6}",
                "read",
                "read_runs",
                "navyfox",
                n,
                bench_read_runs_navyfox,
                ref,
                n,
            )
        finally:
            os.unlink(ref)

    # ---- read: run properties ---------------------------------------------
    print("\n=== Read: run properties (text + bold/italic/color/font_size) ===")
    for n in sizes:
        ref = _build_ref_runs_formatted(n)
        print(f"  [ref n={n} formatted runs in 1 para]")
        try:
            _run(
                f"python-docx  n={n:>6}",
                "read",
                "read_run_props",
                "python_docx",
                n,
                bench_read_run_props_docx,
                ref,
                n,
            )
            _run(
                f"navyfox      n={n:>6}",
                "read",
                "read_run_props",
                "navyfox",
                n,
                bench_read_run_props_navyfox,
                ref,
                n,
            )
        finally:
            os.unlink(ref)

    # ---- read: table cells ------------------------------------------------
    print("\n=== Read: table cells ===")
    print(f"  [capped at max n={table_sizes[-1]}; ref file creation also uses python-docx]")
    for n in table_sizes:
        ref = _build_ref_table(n)
        print(f"  [ref n={n} rows × 3 cols]")
        try:
            _run(
                f"python-docx  n={n:>6}",
                "read",
                "read_table",
                "python_docx",
                n,
                bench_read_table_docx,
                ref,
                n,
            )
            _run(
                f"navyfox      n={n:>6}",
                "read",
                "read_table",
                "navyfox",
                n,
                bench_read_table_navyfox,
                ref,
                n,
            )
        finally:
            os.unlink(ref)

    # ---- read: mixed ------------------------------------------------------
    print("\n=== Read: mixed (paragraphs + tables) ===")
    for n in sizes:
        ref = _build_ref_mixed(n)
        print(f"  [ref n={n}]")
        try:
            _run(
                f"python-docx  n={n:>6}",
                "read",
                "read_mixed",
                "python_docx",
                n,
                bench_read_mixed_docx,
                ref,
                n,
            )
            _run(
                f"navyfox      n={n:>6}",
                "read",
                "read_mixed",
                "navyfox",
                n,
                bench_read_mixed_navyfox,
                ref,
                n,
            )
        finally:
            os.unlink(ref)

    # ---- write CSV --------------------------------------------------------
    out = os.path.join(os.path.dirname(__file__), "benchmark_results.csv")
    fieldnames = ["category", "operation", "library", "n", "repeat", "elapsed_s", "elapsed_ms"]
    with open(out, "w", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(_rows)

    print(f"\nResults written to {out}")
    _print_summary()


def _print_summary() -> None:
    """Print a speedup summary table: navyfox mean / python_docx mean for each (operation, n)."""
    from collections import defaultdict

    # key: (operation, n) → {library: [elapsed_ms, ...]}
    raw: dict[tuple[str, int], dict[str, list[float]]] = defaultdict(lambda: defaultdict(list))
    for row in _rows:
        raw[(str(row["operation"]), int(str(row["n"])))][str(row["library"])].append(
            float(str(row["elapsed_ms"]))
        )

    # key: (operation, n) → {library: mean_ms}
    lookup: dict[tuple[str, int], dict[str, float]] = {
        key: {lib: statistics.mean(vals) for lib, vals in libs.items()} for key, libs in raw.items()
    }

    print("\n=== Speedup summary (python-docx ms / navyfox ms, >1 means navyfox is faster) ===")
    print(f"  {'operation':<22} {'n':>7}  {'docx ms':>10}  {'nfox ms':>10}  {'speedup':>8}")
    print("  " + "-" * 64)
    for (op, n), libs in sorted(lookup.items(), key=lambda kv: (kv[0][0], kv[0][1])):
        if "python_docx" in libs and "navyfox" in libs:
            d = libs["python_docx"]
            nf = libs["navyfox"]
            speedup = d / nf if nf > 0 else float("inf")
            print(f"  {op:<22} {n:>7}  {d:>10.2f}  {nf:>10.2f}  {speedup:>7.2f}×")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="NavyFox vs python-docx benchmark")
    parser.add_argument(
        "--quick",
        action="store_true",
        help="Use smaller sizes for a fast sanity-check run (~5 min vs ~30 min)",
    )
    args = parser.parse_args()
    sizes = SIZES_QUICK if args.quick else SIZES
    table_sizes = TABLE_SIZES_QUICK if args.quick else TABLE_SIZES
    main(sizes, table_sizes)
